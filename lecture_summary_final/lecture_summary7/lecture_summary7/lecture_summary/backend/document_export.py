# Converted from main.ipynb
# -*- coding: utf-8 -*-


# %% In[1]
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
import os
from bidi.algorithm import get_display
import sys

# Ensure console output uses UTF-8 (avoids Windows cp1252 errors)
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

# %% In[2]
import os, re
from dataclasses import dataclass
from typing import List
from pathlib import Path
from html import escape

# Word
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT

# Arabic shaping for PDF
import arabic_reshaper
from bidi.algorithm import get_display

def ar_text(s: str) -> str:
    return escape(get_display(arabic_reshaper.reshape(s or "")))

OUT_DIR = Path("generated_documents")
OUT_DIR.mkdir(parents=True, exist_ok=True)

def pick_arabic_font() -> str:
    candidates = [
        "fonts/Amiri-Regular.ttf",
        "static/fonts/Amiri-Bold.ttf",
        r"C:\Windows\Fonts\trado.ttf",   # Traditional Arabic (Windows)
        r"C:\Windows\Fonts\arial.ttf",   # Arial fallback
        r"/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return ""

AR_FONT_PATH = pick_arabic_font()
print("Using font:", AR_FONT_PATH if AR_FONT_PATH else "(fallback)")

# %% In[3]
@dataclass
class Section:
    heading: str
    body_lines: List[str]

@dataclass
class DocumentModel:
    title: str
    sections: List[Section]

_md_bold_italics = re.compile(r'(\*\*|__|\*|_)(.*?)\1')
_md_heading = re.compile(r'^\s{0,3}(#{1,6})\s+(.*)$')
_md_ul = re.compile(r'^\s*([*\-\+])\s+(.*)$')
_md_ol = re.compile(r'^\s*(\d+)[\.\)]\s+(.*)$')

def _strip_md_inline(s: str) -> str:
    # remove simple **bold**/*italic* markers
    return _md_bold_italics.sub(lambda m: m.group(2), s).strip()

def parse_markdown_to_model(text: str) -> DocumentModel:
    lines = [l.rstrip() for l in text.splitlines() if l.strip() != ""]
    title = "بدون عنوان"
    sections: List[Section] = []

    cur_heading = None
    cur_body: List[str] = []

    # find first H1/H2 as title; treat H2/H3 as section starts
    for i, raw in enumerate(lines):
        m = _md_heading.match(raw)
        if m:
            level = len(m.group(1))
            content = _strip_md_inline(m.group(2))
            if title == "بدون عنوان" and level in (1, 2):
                title = content
                continue
            # new section
            if cur_heading is not None:
                sections.append(Section(cur_heading, cur_body))
            cur_heading = content
            cur_body = []
        else:
            # bullets / numbered / paragraph
            mul = _md_ul.match(raw)
            mol = _md_ol.match(raw)
            if mul:
                cur_body.append("- " + _strip_md_inline(mul.group(2)))
            elif mol:
                # normalize numbered list into bullets
                cur_body.append("- " + _strip_md_inline(mol.group(2)))
            else:
                cur_body.append(_strip_md_inline(raw))

    if cur_heading is not None:
        sections.append(Section(cur_heading, cur_body))

    # if no explicit sections found, put all lines as a single section
    if not sections:
        sections = [Section("المحتوى", lines)]

    return DocumentModel(title=title or "بدون عنوان", sections=sections)

def parse_legacy_to_model(text: str) -> DocumentModel:
    parts = [s.strip() for s in re.split(r'^[ \t]*---[ \t]*$', text, flags=re.MULTILINE)]
    parts = [p for p in parts if p]
    if not parts:
        raise ValueError("النص فارغ أو لا يحتوي قواطع '---'.")

    title_first_line = parts[0].splitlines()
    title = (title_first_line[0] if title_first_line else "").strip() or "بدون عنوان"

    sections: List[Section] = []
    for block in parts[1:]:
        lines = [l.rstrip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        heading = _strip_md_inline(lines[0])
        body = []
        for line in lines[1:]:
            # accept '-', '•', '*', numbered
            if re.match(r'^\s*[\-\u2022\*]\s+', line) or re.match(r'^\s*\d+[\.\)]\s+', line):
                cleaned = re.sub(r'^\s*(?:[\-\u2022\*]|\d+[\.\)])\s+', '', line)
                body.append("- " + _strip_md_inline(cleaned))
            else:
                body.append(_strip_md_inline(line))
        sections.append(Section(heading=heading, body_lines=body))
    return DocumentModel(title=title, sections=sections)

def parse_text_to_model(text: str) -> DocumentModel:
    # Heuristic: if there are Markdown headings (##/###), use MD parser; else try legacy.
    if re.search(r'^\s{0,3}#{2,3}\s+', text, flags=re.MULTILINE):
        return parse_markdown_to_model(text)
    # fallback to legacy
    return parse_legacy_to_model(text)

# %% In[4]
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_paragraph_rtl_and_alignment(paragraph, alignment):
    """
    Sets the paragraph's direction to Right-to-Left (RTL) and its
    alignment to the specified value (e.g., WD_ALIGN_PARAGRAPH.RIGHT).
    This simulates the Ctrl+R shortcut for all text.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # The 'bidi' property is crucial for making text behave like Arabic or Hebrew,
    # ensuring numbers and punctuation are placed correctly in an RTL context.
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

    # Set the alignment
    paragraph.alignment = alignment

def generate_docx_from_model(model: DocumentModel, output_path: str):
   
    doc = Document()

    # Create the main title as a heading (level 1)
    title_heading = doc.add_heading(model.title, level=1)
    # The title should be centered, as requested
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Adds a blank line below the title for spacing

    # Process each section of the document model
    for sec in model.sections:
        # Add a heading for the current section
        section_heading = doc.add_heading(sec.heading, level=2)
        # Apply the RTL and right-alignment formatting to the section heading
        set_paragraph_rtl_and_alignment(section_heading, WD_ALIGN_PARAGRAPH.RIGHT)

        # Process each line of the section's body
        for line in sec.body_lines:
            # Check for bullet points and apply the correct style
            if line.startswith("- "):
                para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                para = doc.add_paragraph(line)

            # Apply the RTL and right-alignment formatting to all body paragraphs
            set_paragraph_rtl_and_alignment(para, WD_ALIGN_PARAGRAPH.RIGHT)

    # Save the final document
    doc.save(output_path)
    print(f"DOCX saved -> {output_path}")

# %% In[5]
def generate_pdf_from_model(model: DocumentModel, output_path: str, font_path: str = AR_FONT_PATH):
    if font_path and os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("AR", font_path))
            font_name = "AR"
        except Exception:
            font_name = "Helvetica"
    else:
        font_name = "Helvetica"  # Arabic may appear disconnected if no Arabic font

    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()

    h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontName=font_name, alignment=TA_RIGHT, spaceAfter=10)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontName=font_name, alignment=TA_RIGHT, spaceAfter=6)
    body = ParagraphStyle('Body', parent=styles['Normal'],   fontName=font_name, alignment=TA_RIGHT, leading=14)
    bullet = ParagraphStyle('Bullet', parent=styles['Normal'], fontName=font_name, alignment=TA_RIGHT, leftIndent=18, bulletIndent=0, leading=14)

    story = []
    story.append(Paragraph(ar_text(model.title), h1))
    story.append(Spacer(1, 12))

    for sec in model.sections:
        story.append(Paragraph(ar_text(sec.heading), h2))
        story.append(Spacer(1, 6))
        for line in sec.body_lines:
            if line.startswith("- "):
                story.append(Paragraph(ar_text("• " + line[2:].strip()), bullet))
            else:
                story.append(Paragraph(ar_text(line), body))
        story.append(Spacer(1, 12))

    doc.build(story)
    print(f"PDF saved  -> {output_path}")

# %% In[6]

INPUT_PATH = "output.txt"

if not os.path.exists(INPUT_PATH):
    raise FileNotFoundError(f"Input file not found: {INPUT_PATH}")

with open(INPUT_PATH, "r", encoding="utf-8") as f:
    raw_text = f.read()

model = parse_text_to_model(raw_text)

docx_path = str(OUT_DIR / "final_document.docx")
pdf_path  = str(OUT_DIR / "final_document.pdf")

generate_docx_from_model(model, docx_path)
generate_pdf_from_model(model, pdf_path)
print("Done.")
